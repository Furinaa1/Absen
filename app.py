import os
import base64
from io import BytesIO
from datetime import datetime, timedelta, timezone
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify, send_from_directory
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.middleware.proxy_fix import ProxyFix
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import qrcode
from openpyxl import Workbook
import cloudinary
import cloudinary.uploader

from models import db, Admin, Peserta, Absensi

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'rahasia-absen-pkl-magang')

# --- KONFIGURASI CLOUDINARY (Membaca dari Environment Variables Vercel) ---
cloudinary.config(
    cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME"),
    api_key=os.environ.get("CLOUDINARY_API_KEY"),
    api_secret=os.environ.get("CLOUDINARY_API_SECRET")
)

# --- KONFIGURASI DATABASE (MENGGUNAKAN CUSTOM_DB_URL & PG8000) ---
raw_db_url = os.environ.get("CUSTOM_DB_URL", "").strip() or os.environ.get("DATABASE_URL", "").strip()

if raw_db_url:
    if raw_db_url.startswith("postgres://"):
        raw_db_url = raw_db_url.replace("postgres://", "postgresql+pg8000://", 1)
    elif raw_db_url.startswith("postgresql://") and not raw_db_url.startswith("postgresql+pg8000://"):
        raw_db_url = raw_db_url.replace("postgresql://", "postgresql+pg8000://", 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = raw_db_url
else:
    if os.environ.get("VERCEL") == "1":
        app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:////tmp/database_absensi.db'
    else:
        app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database_absensi.db'

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Konfigurasi Pool Connection agar koneksi PostgreSQL tahan putus
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    "pool_pre_ping": True,
    "pool_recycle": 300,
}

IS_VERCEL = os.environ.get("VERCEL") == "1"
if IS_VERCEL:
    UPLOAD_FOLDER = os.path.join('/tmp', 'uploads')
else:
    UPLOAD_FOLDER = os.path.join('static', 'uploads')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

db.init_app(app)

# --- SAFE DATABASE INIT FUNCTION ---
def init_db_safe():
    """Menjalankan pembuat tabel dan admin default tanpa menyebabkan crash 500"""
    try:
        db.create_all()
        admin_exists = Admin.query.filter_by(username='instruktur').first()
        if not admin_exists:
            hashed_pw = generate_password_hash('admin1234')
            admin_default = Admin(username='instruktur', password=hashed_pw)
            db.session.add(admin_default)
            db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"Database Init Error (Ignored for safety): {e}")

login_manager = LoginManager()
login_manager.login_view = 'login'
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    try:
        return Admin.query.get(int(user_id))
    except Exception:
        return None

WITA = timezone(timedelta(hours=8))

def get_current_domain():
    """Fungsi helper untuk mendapatkan domain aktif secara dinamis & aman"""
    vercel_url = os.environ.get("VERCEL_URL")
    if vercel_url:
        return f"https://{vercel_url}"
    
    if request.host:
        scheme = request.headers.get('X-Forwarded-Proto', 'https')
        return f"{scheme}://{request.host}"
        
    return "https://absen-4goursy9v-kunii.vercel.app"

# --- HELPER UPLOAD CLOUDINARY YANG DISEMPURNAKAN ---
def upload_to_cloudinary(file_data, folder_name="absensi_pkl"):
    """Mengunggah file (baik base64 string, byte, atau file object) ke Cloudinary dan mengembalikan secure URL."""
    try:
        if not file_data:
            return None
            
        # Jika data berupa string base64 dari JavaScript (canvas/kamera)
        if isinstance(file_data, str) and file_data.startswith('data:image'):
            header, encoded = file_data.split(",", 1)
            file_data = base64.b64decode(encoded)

        response = cloudinary.uploader.upload(file_data, folder=folder_name)
        return response.get("secure_url")
    except Exception as e:
        print(f"Gagal upload ke Cloudinary: {e}")
        return None

# --- ROUTES ---

@app.route('/')
def index():
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    init_db_safe()

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        try:
            admin = Admin.query.filter_by(username=username).first()
            if admin and check_password_hash(admin.password, password):
                login_user(admin)
                return redirect(url_for('admin_dashboard'))
            else:
                flash("Username atau password salah!", "danger")
        except Exception as e:
            flash(f"Terjadi masalah koneksi database: {str(e)}", "danger")
            
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/presensi', methods=['GET', 'POST'])
def presensi():
    if request.method == 'POST':
        nomor_induk = request.form.get('nomor_induk', '').strip()
        status = request.form.get('status', 'Hadir Pagi')
        jurnal = request.form.get('jurnal', '').strip()
        alasan_izin = request.form.get('alasan_izin', '').strip()
        lat = request.form.get('latitude')
        long = request.form.get('longitude')
        foto_base64 = request.form.get('foto_base64')
        ttd_base64 = request.form.get('ttd_base64')

        peserta = Peserta.query.filter_by(nomor_induk=nomor_induk).first()
        if not peserta:
            flash("NIM / NISN tidak terdaftar! Periksa kembali.", "danger")
            return redirect(url_for('presensi'))

        waktu_sekarang = datetime.now(WITA)

        url_foto = None
        if foto_base64:
            url_foto = upload_to_cloudinary(foto_base64, folder_name="selfie_peserta")

        url_ttd = None
        if ttd_base64:
            url_ttd = upload_to_cloudinary(ttd_base64, folder_name="tanda_tangan")

        url_surat = None
        if status not in ['Hadir Pagi', 'Pulang Sore']:
            file_surat = request.files.get('surat_izin')
            if file_surat and file_surat.filename != '':
                url_surat = upload_to_cloudinary(file_surat, folder_name="surat_izin")

        lat_float = None
        long_float = None
        try:
            if lat: lat_float = float(lat)
            if long: long_float = float(long)
        except ValueError:
            pass

        absen_baru = Absensi(
            peserta_id=peserta.id,
            status=status,
            jurnal_harian=jurnal if status == 'Pulang Sore' else None,
            alasan_izin=alasan_izin if status not in ['Hadir Pagi', 'Pulang Sore'] else None,
            foto_selfie=url_foto,
            tanda_tangan=url_ttd,
            surat_izin=url_surat,
            latitude=lat_float,
            longitude=long_float,
            waktu_masuk=waktu_sekarang.replace(tzinfo=None)
        )
        db.session.add(absen_baru)
        db.session.commit()

        flash(f"Pengajuan {status} Berhasil Dikirim untuk {peserta.nama}!", "success")
        return redirect(url_for('presensi'))

    return render_template('presensi.html')

@app.route('/biodata', methods=['GET', 'POST'])
def biodata():
    if request.method == 'POST':
        nomor_induk = request.form.get('nomor_induk', '').strip()
        peserta = Peserta.query.filter_by(nomor_induk=nomor_induk).first()
        if not peserta:
            flash('NIM / NISN belum terdaftar di sistem admin!', 'danger')
            return redirect(url_for('biodata'))
        
        peserta.nama = request.form.get('nama')
        peserta.tempat_lahir = request.form.get('tempat_lahir')
        peserta.tanggal_lahir = request.form.get('tanggal_lahir')
        peserta.jenis_kelamin = request.form.get('jenis_kelamin')
        peserta.hoby = request.form.get('hoby')
        peserta.instansi = request.form.get('instansi')
        peserta.jurusan_kelas = request.form.get('jurusan_kelas')
        peserta.no_hp = request.form.get('no_hp')
        peserta.email = request.form.get('email')
        peserta.alamat = request.form.get('alamat')
        peserta.bidang_penempatan = request.form.get('bidang_penempatan')
        peserta.tanggal_mulai = request.form.get('tanggal_mulai')
        peserta.tanggal_selesai = request.form.get('tanggal_selesai')
        peserta.nama_pembimbing = request.form.get('nama_pembimbing')
        peserta.nip = request.form.get('nip')
        peserta.kategori = request.form.get('kategori', 'PKL')
        
        db.session.commit()
        flash('Biodata lengkap berhasil disimpan!', 'success')
        return redirect(url_for('biodata'))

    return render_template('biodata.html')

@app.route('/admin')
@login_required
def admin_dashboard():
    list_peserta = Peserta.query.all()
    tanggal_mulai = request.args.get('tanggal_mulai')
    tanggal_selesai = request.args.get('tanggal_selesai')
    
    query = Absensi.query
    if tanggal_mulai:
        try:
            dt_mulai = datetime.strptime(tanggal_mulai, '%Y-%m-%d')
            query = query.filter(Absensi.waktu_masuk >= dt_mulai)
        except ValueError:
            pass
            
    if tanggal_selesai:
        try:
            dt_selesai = datetime.strptime(tanggal_selesai + ' 23:59:59', '%Y-%m-%d %H:%M:%S')
            query = query.filter(Absensi.waktu_masuk <= dt_selesai)
        except ValueError:
            pass
            
    rekap_absensi = query.order_by(Absensi.waktu_masuk.desc()).all()
    return render_template('admin_dashboard.html', peserta=list_peserta, absensi=rekap_absensi, tanggal_mulai=tanggal_mulai, tanggal_selesai=tanggal_selesai)

@app.route('/admin/biodata-peserta')
@login_required
def admin_biodata_peserta():
    all_peserta = Peserta.query.all()
    return render_template('admin_biodata_peserta.html', peserta_list=all_peserta)

@app.route('/admin/peserta')
@login_required
def admin_peserta():
    all_peserta = Peserta.query.all()
    return render_template('admin_peserta.html', peserta_list=all_peserta)

@app.route('/admin/peserta/cetak/<int:id>')
@login_required
def cetak_biodata_peserta(id):
    peserta = Peserta.query.get_or_404(id)
    return render_template('cetak_biodata.html', peserta=peserta)

@app.route('/admin/peserta/tambah', methods=['POST'])
@login_required
def tambah_peserta():
    nama = request.form.get('nama')
    nomor_induk = request.form.get('nomor_induk', '').strip()
    kategori = request.form.get('kategori')
    instansi = request.form.get('instansi')
    
    if Peserta.query.filter_by(nomor_induk=nomor_induk).first():
        flash("NIM/NISN sudah terdaftar!", "warning")
    else:
        peserta_baru = Peserta(nama=nama, nomor_induk=nomor_induk, kategori=kategori, instansi=instansi)
        db.session.add(peserta_baru)
        db.session.commit()
        flash("Peserta berhasil ditambahkan!", "success")
        
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/peserta/hapus/<int:id>')
@login_required
def hapus_peserta(id):
    peserta = Peserta.query.get_or_404(id)
    absensi_records = Absensi.query.filter_by(peserta_id=peserta.id).all()
    
    for absensi in absensi_records:
        db.session.delete(absensi)

    db.session.delete(peserta)
    db.session.commit()
    flash("Data peserta beserta absensinya berhasil dihapus!", "success")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/absensi/hapus/<int:id>')
@login_required
def hapus_absensi(id):
    absensi = Absensi.query.get_or_404(id)
    db.session.delete(absensi)
    db.session.commit()
    flash("Data riwayat absensi berhasil dihapus!", "success")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/export-word-biodata')
@login_required
def export_word_biodata():
    doc = Document()
    list_peserta = Peserta.query.all()
    if not list_peserta:
        doc.add_heading('Rekap Biodata Lengkap Peserta PKL/Magang', level=1)
        doc.add_paragraph('Belum ada data peserta.')
    else:
        for idx, p in enumerate(list_peserta, 1):
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run('Rekap Biodata Lengkap Peserta PKL/Magang\n')
            run_title.bold = True
            run_title.font.size = Pt(14)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Kolom Data'
            hdr_cells[1].text = 'Keterangan'
            
            data_fields = [
                ("Nama Lengkap", p.nama), ("NIM / NISN", p.nomor_induk),
                ("Kategori", p.kategori), ("Tempat Lahir", p.tempat_lahir),
                ("Tanggal Lahir", p.tanggal_lahir), ("Jenis Kelamin", p.jenis_kelamin),
                ("Hobi", p.hoby), ("Asal Sekolah / Kampus", p.instansi),
                ("Jurusan / Kelas", p.jurusan_kelas), ("No HP", p.no_hp),
                ("Email", p.email), ("Alamat", p.alamat),
                ("Bidang Penempatan", p.bidang_penempatan), ("Tanggal Mulai", p.tanggal_mulai),
                ("Tanggal Selesai", p.tanggal_selesai), ("Nama Pembimbing", p.nama_pembimbing),
                ("NIP Pembimbing", p.nip)
            ]
            
            for label, val in data_fields:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(val if val else '-')
                
            doc.add_paragraph()
            ttd_table = doc.add_table(rows=1, cols=2)
            cell_kiri, cell_kanan = ttd_table.rows[0].cells[0], ttd_table.rows[0].cells[1]

            p_foto = cell_kiri.paragraphs[0]
            p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_foto.add_run("Pas foto\n3 x 4\nBerwarna")

            p_ttd = cell_kanan.paragraphs[0]
            p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_ttd.add_run("Samarinda, ...................................\n")
            p_ttd.add_run("Hormat saya\n\n\n\n")
            run_ttd_nama = p_ttd.add_run(f"( {p.nama if p.nama else '...................................'} )")
            run_ttd_nama.bold = True

            if idx < len(list_peserta):
                doc.add_page_break()

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='Rekap_Biodata_Peserta.docx'
    )

@app.route('/export_excel')
@login_required
def export_excel():
    list_peserta = Peserta.query.all()
    
    wb = Workbook()
    default_sheet = wb.active
    wb.remove(default_sheet)
    
    if not list_peserta:
        ws = wb.create_sheet(title='Tidak Ada Peserta')
        ws.append(['Waktu Masuk', 'Nama Peserta', 'Nomor Induk', 'Kategori', 'Instansi', 'Status', 'Keterangan', 'GPS', 'Foto', 'TTD'])
    else:
        for p in list_peserta:
            nama_sheet = "".join(c for c in (p.nama or "") if c.isalnum() or c in (' ', '_'))[:30].strip()
            if not nama_sheet:
                nama_sheet = f"Peserta_{p.id}"
            
            ws = wb.create_sheet(title=nama_sheet)
            headers = ['Waktu Masuk', 'Nama Peserta', 'Nomor Induk', 'Kategori', 'Instansi', 'Status', 'Keterangan / Jurnal / Alasan', 'Koordinat GPS', 'Foto Selfie', 'Tanda Tangan']
            ws.append(headers)
            
            ws.column_dimensions['I'].width = 35
            ws.column_dimensions['J'].width = 35
            
            data_absensi_peserta = Absensi.query.filter_by(peserta_id=p.id).order_by(Absensi.waktu_masuk.desc()).all()
            
            for a in data_absensi_peserta:
                waktu_str = a.waktu_masuk.strftime('%Y-%m-%d %H:%M:%S') if a.waktu_masuk else ''
                keterangan = a.jurnal_harian if a.status == 'Pulang Sore' else a.alasan_izin
                koordinat = f"{a.latitude}, {a.longitude}" if a.latitude and a.longitude else ''
                
                link_foto = f'=HYPERLINK("{a.foto_selfie}", "Lihat Foto")' if a.foto_selfie else '-'
                link_ttd = f'=HYPERLINK("{a.tanda_tangan}", "Lihat TTD")' if a.tanda_tangan else '-'
                
                row_data = [waktu_str, p.nama or '', p.nomor_induk or '', p.kategori or '', p.instansi or '', a.status, keterangan or '', koordinat, link_foto, link_ttd]
                ws.append(row_data)

    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='rekap_absensi_per_peserta.xlsx'
    )

@app.route('/api/get-nama/<nomor_induk>')
def get_nama(nomor_induk):
    peserta = Peserta.query.filter_by(nomor_induk=nomor_induk).first()
    if peserta:
        return jsonify({"success": True, "nama": peserta.nama, "kategori": peserta.kategori})
    else:
        return jsonify({"success": False, "message": "NIM / NISN tidak ditemukan!"})

@app.route('/admin/generate-qr')
@login_required
def generate_qr():
    domain_aktif = get_current_domain()
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    
    url_presensi = f"{domain_aktif}/presensi"
    qr_presensi = qrcode.make(url_presensi)
    qr_presensi.save(os.path.join(UPLOAD_FOLDER, 'qr_absensi.png'))

    url_biodata = f"{domain_aktif}/biodata"
    qr_biodata = qrcode.make(url_biodata)
    qr_biodata.save(os.path.join(UPLOAD_FOLDER, 'qr_biodata.png'))

    flash(f"QR Code Presensi & Biodata berhasil diperbarui menggunakan domain: {domain_aktif}", "success")
    return redirect(url_for('admin_dashboard'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)